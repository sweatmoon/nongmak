"""
신고서 DOCX 생성 모듈
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io


DISCLAIMER = "※ 본 문서는 제출용 초안이며, 최종 제출 전 관할 지자체 및 등록 시공업체의 검토가 반드시 필요합니다."


def _set_font(run, name="맑은 고딕", size=10, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), name)
    rPr.insert(0, rFonts)


def _heading(doc, text, level=1, color=(30, 64, 175)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    _set_font(run, size=13 if level == 1 else 11, bold=True, color=color)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p


def _add_table_row(table, label, value, label_color=(239, 246, 255)):
    row = table.add_row()
    row.cells[0].text = label
    row.cells[1].text = value
    for i, cell in enumerate(row.cells):
        for para in cell.paragraphs:
            for run in para.runs:
                _set_font(run, size=9, bold=(i == 0))
        cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        if i == 0:
            shd.set(qn("w:fill"), "EFF6FF")
        else:
            shd.set(qn("w:fill"), "FFFFFF")
        cell._tc.get_or_add_tcPr().append(shd)


def _add_border(table):
    """테이블 테두리 설정"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "CBD5E1")
        tblBorders.append(border)
    tblPr.append(tblBorders)


def generate_septic_docx(order_id: str, req, computed: dict) -> bytes:
    """개인하수처리시설 신고서 초안 생성"""
    doc = Document()

    # 페이지 여백
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 제목
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("개인하수처리시설 설치 신고서 (초안)")
    _set_font(run, size=16, bold=True, color=(30, 64, 175))
    title_p.paragraph_format.space_after = Pt(4)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle_p.add_run("※ 본 문서는 시스템 자동 생성 초안입니다. 관할 지자체 공식 서식으로 재작성 후 제출하시기 바랍니다.")
    _set_font(run2, size=9, color=(220, 38, 38))
    subtitle_p.paragraph_format.space_after = Pt(12)

    # 신청인 정보
    _heading(doc, "1. 신청인(설치자) 정보")
    t1 = doc.add_table(rows=0, cols=2)
    t1.style = "Table Grid"
    t1.columns[0].width = Cm(5)
    t1.columns[1].width = Cm(12)
    _add_table_row(t1, "성명(법인명)", req.applicant_name)
    _add_table_row(t1, "연락처", req.phone or "")
    _add_table_row(t1, "이메일", req.email or "")
    _add_border(t1)

    doc.add_paragraph()

    # 시설 정보
    _heading(doc, "2. 시설 정보")
    t2 = doc.add_table(rows=0, cols=2)
    t2.style = "Table Grid"
    t2.columns[0].width = Cm(5)
    t2.columns[1].width = Cm(12)
    _add_table_row(t2, "설치 장소(주소)", req.address)
    _add_table_row(t2, "시설 종류", "개인하수처리시설 (정화조)")
    _add_table_row(t2, "처리 용량", f"{computed['septic_capacity_m3']} m³/일")
    _add_table_row(t2, "처리 방식", _get_treatment_mode_label(req.treatment_mode.value))
    _add_table_row(t2, "사용 인원", f"상시 {req.occupants_regular}명 / 최대 {req.occupants_max}명")
    _add_table_row(t2, "화장실 유형", _get_toilet_type_label(req.toilet_type.value))
    _add_border(t2)

    doc.add_paragraph()

    # 시공 예정 정보
    _heading(doc, "3. 시공 예정")
    t3 = doc.add_table(rows=0, cols=2)
    t3.style = "Table Grid"
    t3.columns[0].width = Cm(5)
    t3.columns[1].width = Cm(12)
    _add_table_row(t3, "시공 예정일", "                년    월    일")
    _add_table_row(t3, "준공 예정일", "                년    월    일")
    _add_table_row(t3, "시공업체 상호", "")
    _add_table_row(t3, "시공업체 연락처", "")
    _add_table_row(t3, "시공업체 등록번호", "")
    _add_border(t3)

    doc.add_paragraph()

    # 첨부 서류
    _heading(doc, "4. 첨부 서류 확인")
    attachments = [
        ("□ 배치도", "정화조 및 농막 위치 표시 (본 패키지 01_배치도.pdf 참고)"),
        ("□ 배관도", "배관 경로 및 구경 표시 (본 패키지 02_배관도.pdf 참고)"),
        ("□ 용량산정서", "(본 패키지 03_용량산정서.pdf 참고, 시공업체 확인 필요)"),
        ("□ 정화조 제품 사양서/형식승인서", "시공업체 제공"),
        ("□ 시공업체 등록증 사본", "전문업체 등록 확인"),
        ("□ 토지 등기부등본", "최근 3개월 이내"),
        ("□ 기타 지자체 요청 서류", "담당 부서 문의"),
    ]
    for label, note in attachments:
        p = doc.add_paragraph()
        run_label = p.add_run(f"{label}  ")
        _set_font(run_label, size=10, bold=True)
        run_note = p.add_run(f"({note})")
        _set_font(run_note, size=9, color=(100, 116, 139))

    doc.add_paragraph()

    # 신고 문구
    sig_p = doc.add_paragraph()
    sig_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sig = sig_p.add_run(
        "위와 같이 개인하수처리시설의 설치를 신고합니다.\n\n"
        "                년    월    일\n\n"
        "신고인(설치자):                                (서명 또는 인)\n\n"
        "시장·군수·구청장 귀하"
    )
    _set_font(run_sig, size=10)
    sig_p.paragraph_format.space_before = Pt(12)

    doc.add_paragraph()

    # 면책 문구
    disc_p = doc.add_paragraph()
    disc_run = disc_p.add_run(DISCLAIMER)
    _set_font(disc_run, size=8, color=(100, 116, 139))
    disc_p.paragraph_format.space_before = Pt(12)

    note_p = doc.add_paragraph()
    note_run = note_p.add_run(f"[자동 생성 초안] 주문번호: {order_id} | 룰셋: {computed['ruleset_id']}")
    _set_font(note_run, size=8, color=(148, 163, 184))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_hut_docx(order_id: str, req, computed: dict) -> bytes:
    """가설건축물 신고서 초안 생성"""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 제목
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("가설건축물 축조 신고서 (초안)")
    _set_font(run, size=16, bold=True, color=(30, 64, 175))
    title_p.paragraph_format.space_after = Pt(4)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle_p.add_run("※ 본 문서는 시스템 자동 생성 초안입니다. 관할 지자체 공식 서식으로 재작성 후 제출하시기 바랍니다.")
    _set_font(run2, size=9, color=(220, 38, 38))
    subtitle_p.paragraph_format.space_after = Pt(12)

    # 신청인 정보
    _heading(doc, "1. 신청인 정보")
    t1 = doc.add_table(rows=0, cols=2)
    t1.style = "Table Grid"
    t1.columns[0].width = Cm(5)
    t1.columns[1].width = Cm(12)
    _add_table_row(t1, "성명(대표자)", req.applicant_name)
    _add_table_row(t1, "연락처", req.phone or "")
    _add_table_row(t1, "이메일", req.email or "")
    _add_border(t1)

    doc.add_paragraph()

    # 건축물 정보
    _heading(doc, "2. 건축물 정보")
    t2 = doc.add_table(rows=0, cols=2)
    t2.style = "Table Grid"
    t2.columns[0].width = Cm(5)
    t2.columns[1].width = Cm(12)
    _add_table_row(t2, "대지 위치(지번)", req.address)
    _add_table_row(t2, "용도", "농막(농업용 가설건축물)")
    _add_table_row(t2, "구조", "경량 철골조 또는 목조 (시공업체 확인)")
    _add_table_row(t2, "건축 면적", f"{req.hut_area_m2} ㎡")
    _add_table_row(t2, "연면적", f"{req.hut_area_m2} ㎡ (1층)")
    _add_table_row(t2, "규모", f"1층 / 가로 {req.hut_w_m}m × 세로 {req.hut_d_m}m")
    _add_table_row(t2, "층수", "지상 1층")
    _add_table_row(t2, "높이", "약 3m 이하 (현장 확인)")
    _add_border(t2)

    doc.add_paragraph()

    # 설치 기간
    _heading(doc, "3. 설치 기간")
    t3 = doc.add_table(rows=0, cols=2)
    t3.style = "Table Grid"
    t3.columns[0].width = Cm(5)
    t3.columns[1].width = Cm(12)
    _add_table_row(t3, "설치 예정일", "                년    월    일")
    _add_table_row(t3, "존치 기간", "3년 (연장 신청 가능, 농지법 확인)")
    _add_table_row(t3, "철거 예정일", "설치일로부터 3년 이내")
    _add_border(t3)

    doc.add_paragraph()

    # 첨부 서류
    _heading(doc, "4. 첨부 서류 확인")
    attachments = [
        ("□ 배치도", "농막 위치 및 형태 표시 (본 패키지 01_배치도.pdf 참고)"),
        ("□ 평면도", "농막 내부 배치도 (시공업체 작성 권장)"),
        ("□ 입면도", "(시공업체 제공 또는 제조사 규격서 활용)"),
        ("□ 토지 등기부등본", "최근 3개월 이내"),
        ("□ 토지이용계획확인서", "토지 용도지역 확인"),
        ("□ 농지원부 또는 영농확인서", "농지법 적용 여부 확인"),
        ("□ 인감증명서 또는 본인서명사실확인서", ""),
        ("□ 기타 지자체 요청 서류", "담당 부서 문의"),
    ]
    for label, note in attachments:
        p = doc.add_paragraph()
        run_label = p.add_run(f"{label}  ")
        _set_font(run_label, size=10, bold=True)
        if note:
            run_note = p.add_run(f"({note})")
            _set_font(run_note, size=9, color=(100, 116, 139))

    doc.add_paragraph()

    # 주의사항
    _heading(doc, "5. 농막 설치 주의사항", color=(220, 38, 38))
    notices = [
        "농막은 농지법에 따라 농업용으로만 사용하여야 하며, 주거 목적으로 사용 불가합니다.",
        "건축면적 20㎡(약 6평) 초과 또는 연면적 33㎡ 초과 시 건축허가 대상이 될 수 있습니다.",
        "관할 지자체별 조례 및 기준이 다를 수 있으므로 반드시 사전 상담이 필요합니다.",
        "농막 내 취사·숙박 시설 설치 여부는 지자체별 허용 기준을 확인하시기 바랍니다.",
    ]
    for n in notices:
        p = doc.add_paragraph()
        run_n = p.add_run(f"• {n}")
        _set_font(run_n, size=9, color=(100, 116, 139))

    doc.add_paragraph()

    # 신고 문구
    sig_p = doc.add_paragraph()
    sig_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sig = sig_p.add_run(
        "위와 같이 가설건축물의 축조를 신고합니다.\n\n"
        "                년    월    일\n\n"
        "신고인:                                (서명 또는 인)\n\n"
        "시장·군수·구청장 귀하"
    )
    _set_font(run_sig, size=10)
    sig_p.paragraph_format.space_before = Pt(12)

    doc.add_paragraph()

    # 면책 문구
    disc_p = doc.add_paragraph()
    disc_run = disc_p.add_run(DISCLAIMER)
    _set_font(disc_run, size=8, color=(100, 116, 139))
    disc_p.paragraph_format.space_before = Pt(12)

    note_p = doc.add_paragraph()
    note_run = note_p.add_run(f"[자동 생성 초안] 주문번호: {order_id} | 룰셋: {computed['ruleset_id']}")
    _set_font(note_run, size=8, color=(148, 163, 184))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _get_treatment_mode_label(mode: str) -> str:
    return {
        "SEPTIC_DISCHARGE": "방류형 (처리수 방류)",
        "INFILTRATION": "침투형 (지중 침투)",
        "UNKNOWN": "미확정 (시공 전 확정 필요)",
    }.get(mode, mode)


def _get_toilet_type_label(t: str) -> str:
    return {
        "FLUSH": "수세식",
        "PORTABLE": "이동식",
        "HOLDING_TANK": "저장조식",
    }.get(t, t)
